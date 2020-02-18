import pandas as pd
import docx
import os
from tqdm import tqdm
from gensim import corpora, models
from collections import OrderedDict
import seaborn as sns
from nltk.corpus import stopwords
from string import punctuation
import re
import pyLDAvis.gensim
import matplotlib.pyplot as plt


def get_docx_files(folder):
    """
    look at the folder with documents and search for docx files
    """
    doc_list = []
    docx_list = []

    for path, subdirs, files in os.walk(folder):
        for name in files:
            if os.path.splitext(os.path.join(path, name))[1] == ".docx":
                doc_list.append(os.path.join(path, name))

    for doc_path in doc_list:
        document = docx.Document(doc_path)
        docx_list.append(document)
    return docx_list


def generate_corpus(txt_file, docx_list):
    """
    generate a corpus from the docx files and exports as a txt file
    """
    corpus = []
    for i in tqdm(docx_list):
        for j in range(len(i.paragraphs)):
            text = i.paragraphs[j].text
            corpus.append(text)

    with open(txt_file, 'w') as file:
        for listitem in corpus:
            file.write('%s\n' % listitem)
    return corpus


def load_corpus_file(txt_file):
    """
    load a txt corpus file, if it already exists
    """
    # LOAD CORPUS FILE
    # open file and read the content in a list
    corpus = []
    with open(txt_file, 'r') as filehandle:
        for line in filehandle:
            # remove linebreak which is the last character of the string
            currentplace = line[:-1]
            # add item to the list
            corpus.append(currentplace)
    return corpus


def replaces_in_corpus(corpus):
    """
    manual term replacements in corpus
    """
    corpus = [i.lower().replace("partes relacionadas ", "partes_relacionadas ") for i in corpus]
    corpus = [i.lower().replace("parte relacionada ", "partes_relacionadas ") for i in corpus]
    corpus = [i.lower().replace("fonte:", "fontes") for i in corpus]
    corpus = [i.lower().replace("contratação", "contratações") for i in corpus]
    corpus = [i.lower().replace("contrato ", "contratações") for i in corpus]
    corpus = [i.lower().replace("contratos", "contratações") for i in corpus]
    corpus = [i.lower().replace("contratoss", "contratações") for i in corpus]
    corpus = [i.lower().replace("princípio", "princípios") for i in corpus]
    corpus = [i.lower().replace("empresa ", "empresas") for i in corpus]
    corpus = [i.lower().replace("acionista ", "acionistas") for i in corpus]
    corpus = [i.lower().replace("valor ", "valores") for i in corpus]
    corpus = [i.lower().replace("aeroporto ", "aeroportos ") for i in corpus]
    corpus = [i.lower().replace("mercado.", "mercado") for i in corpus]
    corpus = [i.lower().replace("mercado,", "mercado") for i in corpus]
    corpus = [i.lower().replace("concessionária ", "concessionárias ") for i in corpus]
    corpus = [i.lower().replace("bndespar ", "bndes ") for i in corpus]
    return corpus


blacklist = ['_','–',' ','Augusto','augusto','Carlos' "$",")","(",'/2015',"walton",'alencar','alves',".","(...)","[...]", 'sobre', 'sr.', 'art.', 'que', 'ser', 'ii', 'ii,', 'inciso,', 'de$', 'nº',
            '(peça','§','(cpf','tc','maria','josé','costa','que,','quanto','(',')','75.','19.','-.','(peças-11),','silva', 'tcu', 'ainda','além', 'qualquer','parte','partes','item','tais','dias','tal',
            'ano','cada']


def remove_blacklist(corpus, blacklist):
    """
    removes a previously defined blacklist from the corpus, and also removes punctuation and shortwords
    """
    shortword = re.compile(r'\W*\b\w{1,1}\b') # # regex to remove word up to 3 chars (inclusive)
    stoplist = stopwords.words('portuguese') + list(punctuation) + blacklist
    stoplist.remove('não')
    stoplist.remove('sem')
    text_total = [[i for i in str(shortword.sub('', str(j))).lower().split() if i not in stoplist] for j in corpus]
    return text_total


def filter_text_chunks(texts, topic):
    """
    applies user defined hard coded filter to pop out words that doesnt help the analysis (domain knowledge)
    """
    filt_part1 = ['partes_relacionadas']
    filt_part2 = ['parte_relacionada']
    filt_part_add = ['participação', 'contratos ', 'contratos', 'contratação ', 'transações', 'contratos', 'dirigente',
                     'termos e condições de mercado', 'contratos', 'notas explicativas', 'contratação ', 'subcontratação',
                     'conflito de interesse', 'contratação ', 'contratos', 'contratos', 'acionistas', 'contratos',
                     'termos e condições de mercado', 'subcontratação', 'transparência ', 'governança', 'contratações',
                     'fora das condições de mercado', 'contratações', 'riscos', 'coligadas', 'coligadas', 'risco', 'contratações',
                     'contratos', 'acionistas', 'contratos', 'contoladas', 'parâmetro de mercado', 'acionistas', 'governança corporativa',
                     'empresas subsidárias', 'condições de mercado', 'acionistas', 'favorecimento indevido', 'condições diferentes',
                     'obras', 'serviços de engenharia', 'favorecer', 'sócio privado', 'arm´s length', 'relacionamentos',
                     'grau de parentesco', 'preços superfaturados', 'beneficiar', 'influenciar', 'sócio', 'contratação ',
                     'sócios comuns', 'acionistas', 'contratos', 'acionistas', 'contratos', 'favorecimento', 'sócios em comum', 'contratações',
                     'contratos', 'notas explicativas', 'contratação ', 'contratos', 'contratação ', 'subcontratação', 'notas explicativas',
                     'condições de mercado', 'condições de mercado', 'fornecedores', 'acionistas', 'contratos', 'termos',
                     'condições diferentes das praticadas no mercado', 'riscos', 'contratos', 'termos e condições de mercado',
                     'diferentes das praticadas no mercado']

    filt_nep = ['nepotismo']
    filt_nep_add = ['cargo em comissão', 'contratação', 'contratado', 'vício', 'contratação de parentes, consanguineos ou por afinidade',
                    'contratados irregularmente', 'favorecimento de parentes', 'apadrinhamento', 'ausência de processo seletivo', 'irmã',
                    'primo', 'tio', 'proibição', 'vedação', 'violados', 'princíipios constitucionais', 'conflito de interesses', 'transnepotismo',
                    'moralidade administrativa', 'nomeação em cargo de comissão', 'cunhada', 'noiva', 'relação']

    filt_mult = ['multas']
    filt_mult_add = ['multa', '$', 'penalidade', 'punição']
    text_term=[]
    if topic == 'PARTES':
        text_term = [j for j in texts if all(i in j for i in filt_part1 or filt_part2) and any(i in j for i in filt_part_add)]
    elif topic == 'NEPOTISMO':
        text_term = [j for j in texts if all(i in j for i in filt_nep) and any(i in j for i in filt_nep_add)]
    elif topic == 'MULTAS':
        text_term = [j for j in texts if all(i in j for i in filt_mult) and any(i in j for i in filt_mult_add)]
    return text_term


def filter_corpus(text_term, text_total, which_text='TOTAL'):
    """
    defines if generated corpus is based on whole data or only on filtered terms data
    """
    if which_text == 'TERM':
        dict_term = corpora.Dictionary(text_term)
        corpus_term = [dict_term.doc2bow(i) for i in text_term]
        tfidf_term = models.TfidfModel(corpus_term)  # step 1 -- initialize a model
        corpus_tfidf = tfidf_term[corpus_term]  # step 2 -- use the model to transform vectors
    elif which_text == 'TOTAL':
        dict_total = corpora.Dictionary(text_total)
        corpus_total = [dict_total.doc2bow(i) for i in text_total]
        tfidf_total = models.TfidfModel(corpus_total)  # step 1 -- initialize a model
        corpus_tfidf = tfidf_total[corpus_total]  # step 2 -- use the model to transform vectors
    return corpus_tfidf


def generate_lda(total_topics, which_text='TOTAL'):
    """
    LDA corpus generation based on choice of whole corpus data or filtered corpus data
    """
    if which_text == 'TERM':
        dict_term = corpora.Dictionary(text_term)
        corpus_term = [dict_term.doc2bow(i) for i in text_term]
        lda_term = models.LdaModel(corpus_term, id2word=dict_term, num_topics=total_topics)
        corpus_lda_term = lda_term[corpus_term]  # create a double wrapper over the original corpus: bow->tfidf->fold-in-lsi\
        corpus_out = corpus_lda_term
        out = [lda_term, corpus_out]
    elif which_text == 'TOTAL':
        dict_total = corpora.Dictionary(text_total)
        corpus_total = [dict_total.doc2bow(i) for i in text_total]
        lda_total = models.LdaModel(corpus_total, id2word=dict_total, num_topics=total_topics)
        corpus_lda_total = lda_total[corpus_total]  # create a double wrapper over the original corpus: bow->tfidf->fold-in-lsi
        corpus_out = corpus_lda_total
        out = [lda_term, corpus_lda_total]
    return out


def show_N_terms(lda_term, lda_total, total_topics):
    """
    show terms by topic, in relevance order
    """
    # Show important words in the topics:
    print(lda_term.show_topics(total_topics))
    print(lda_total.show_topics(total_topics))
    return


def prepare_data_clustermap(lda_term, lda_total, total_topics):
    """
    generates input dataframes for posterior clustermap generation
    """
    data_lda_term = {i: OrderedDict(lda_term.show_topic(i, 10)) for i in range(total_topics)}
    data_lda_total = {i: OrderedDict(lda_total.show_topic(i, 10)) for i in range(total_topics)}

    df_lda_term = pd.DataFrame(data_lda_term).fillna(0)
    df_lda_term = df_lda_term.T
    print("df_lda_term shape:", df_lda_term.shape)

    df_lda_total = pd.DataFrame(data_lda_total).fillna(0)
    df_lda_total = df_lda_total.T
    print("df_lda_total shape:", df_lda_total.shape)
    return df_lda_term, df_lda_total


def generate_clustermap(which_text='TOTAL'):
    """
    clustermap generation, base on whole data or on filtered corpus data
    """
    if  which_text='TERM':
        g = sns.clustermap(df_lda_term.corr(), standard_scale=1, center=0, cmap="RdBu",
                           metric='cosine', linewidths=.05, figsize=(15, 15))
        plt.setp(g.ax_heatmap.yaxis.get_majorticklabels(), rotation=0)
        # plt.savefig('nepotismo_filtro_sem_arquivo_grande.jpg')
        plt.show()

    if  which_text='TOTAL':
        g = sns.clustermap(df_lda_total.corr(), standard_scale=1, center=0, cmap="RdBu",
                           metric='cosine', linewidths=.05, figsize=(15, 15))
        plt.setp(g.ax_heatmap.yaxis.get_majorticklabels(), rotation=0)
        # plt.savefig('nepotismo_total_sem_arquivo_grande.jpg')
        plt.show()
    return


def generate_lda_vis(which_text='TOTAL'):
    """
    LDA model visualization
    """
    if which_text='TOTAL':
        pyLDAvis.enable_notebook()
        panel_tot = pyLDAvis.gensim.prepare(lda_total, corpus_total, dict_total, mds='PCoA')
        panel_tot
    elif which_text='TERM'
        pyLDAvis.enable_notebook()
        panel = pyLDAvis.gensim.prepare(lda_term, corpus_lda_term, dict_term, mds='PCoA')
        panel
    return