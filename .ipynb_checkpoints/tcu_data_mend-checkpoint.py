# -*- coding: utf-8 -*-

import pandas as pd
import docx
import os
from tqdm import tqdm
from gensim import corpora, models
from collections import OrderedDict
import seaborn as sns
from nltk.corpus import stopwords
from string import punctuation
import re
import pyLDAvis.gensim
import matplotlib.pyplot as plt
import nltk
nltk.download('stopwords')


def get_docx_files(folder):
    """
    look at the folder with documents and search for docx files
    """
    doc_list = []
    docx_list = []

    for path, subdirs, files in os.walk(folder):
        for name in files:
            if os.path.splitext(os.path.join(path, name))[1] == ".docx":
                doc_list.append(os.path.join(path, name))

    for doc_path in doc_list:
        document = docx.Document(doc_path)
        docx_list.append(document)
    return docx_list


def generate_corpus(txt_output_file, docx_list):
    """
    generates a corpus from the docx files and export it as a txt file
    """
    corpus = []
    for i in tqdm(docx_list):
        for j in range(len(i.paragraphs)):
            text = i.paragraphs[j].text
            corpus.append(text)

    with open(txt_output_file, 'w') as file:
        for listitem in corpus:
            file.write('%s\n' % listitem)
    return corpus


def load_corpus_file(txt_file):
    """
    load a txt corpus file, if it already exists
    """
    # LOAD CORPUS FILE
    # open file and read the content in a list
    corpus = []
    with open(txt_file, 'r', encoding='utf-8') as filehandle:
        for line in filehandle:
            # remove linebreak which is the last character of the string
            currentplace = line[:-1]
            # add item to the list
            corpus.append(currentplace)
    return corpus


def replaces_in_corpus(corpus, topic='PARTES'):
    """
    manual term replacements in corpus
    """
    corpus = [i.lower() for i in corpus]
    if topic == 'PARTES':
        corpus = [i.replace("partes relacionadas ", "partes_relacionadas ") for i in corpus]
        corpus = [i.replace("parte relacionada ", "partes_relacionadas ") for i in corpus]
        corpus = [i.replace("fonte:", "fontes") for i in corpus]
        corpus = [i.replace("contratação", "contratações") for i in corpus]
        corpus = [i.replace("contrato ", "contratações") for i in corpus]
        corpus = [i.replace("contratos", "contratações") for i in corpus]
        corpus = [i.replace("contratoss", "contratações") for i in corpus]
        corpus = [i.replace("princípio", "princípios") for i in corpus]
        corpus = [i.replace("empresa ", "empresas") for i in corpus]
        corpus = [i.replace("acionista ", "acionistas") for i in corpus]
        corpus = [i.replace("valor ", "valores") for i in corpus]
        corpus = [i.replace("aeroporto ", "aeroportos ") for i in corpus]
        corpus = [i.replace("mercado.", "mercado") for i in corpus]
        corpus = [i.replace("mercado,", "mercado") for i in corpus]
        corpus = [i.replace("concessionária ", "concessionárias ") for i in corpus]
        # corpus = [i.replace("bndespar ","bndes ") for i in corpus] # uncomment this row only if row 165 is used instead of 166 (i.e. filt_part_add list being used for filtering)
    elif topic == 'NEPOTISMO':
        corpus = [i.replace("partes relacionadas", "partes_relacionadas") for i in corpus]
        corpus = [i.replace("parte relacionada", "parte_relacionada") for i in corpus]
        corpus = [i.replace("ato ", "atos") for i in corpus]
        corpus = [i.replace("valor ", "valores") for i in corpus]
        corpus = [i.replace("processo ", "processos") for i in corpus]
        corpus = [i.replace("cargo ", "cargos") for i in corpus]
        corpus = [i.replace("princípio ", "princípios ") for i in corpus]
        corpus = [i.replace("irregularidade ", "irregularidades ") for i in corpus]
    elif topic == 'CONF_INTERESSE':
        corpus = [i.replace("contrato ", "contratos") for i in corpus]
        corpus = [i.replace("empresa ", "empresas") for i in corpus]
        corpus = [i.replace("valor ", "valores") for i in corpus]
        corpus = [i.replace("público,", "público") for i in corpus]
        corpus = [i.replace("público.", "público") for i in corpus]
        corpus = [i.replace("pública,", "público") for i in corpus]
        corpus = [i.replace("projeto,", "projetos") for i in corpus]
        corpus = [i.replace("princípio,", "princípios") for i in corpus]
        corpus = [re.sub(r'\bconflito de interesse\b', 'conflito_de_interesses', i) for i in corpus]
        corpus = [re.sub(r'\bconflito de interesses\b', 'conflito_de_interesses', i) for i in corpus]
        corpus = [re.sub(r'\bconflitos de interesses\b', 'conflito_de_interesses', i) for i in corpus]
        corpus = [re.sub(r'\binteresses\b', 'conflito_de_interesses', i) for i in corpus]
        corpus = [re.sub(r'\binteresse\b', 'conflito_de_interesses', i) for i in corpus]
    return corpus


blacklist_conflito = ['_','–',' ','Augusto','augusto','Carlos' "$",")","(",'/2015',"walton",'alencar','alves',".","(...)",
             "[...]", 'sobre', 'sob', 'sr.', 'art.', 'que', 'ser', 'ii', 'ii,', 'inciso,', 'de$', 'nº','(peça','§',
             '(cpf','tc','maria','josé','costa','que,','quanto','(',')','75.','19.','-.','(peças-11),',
             'silva', 'tcu', 'ainda','além', 'qualquer','parte','partes','item','tais','dias','tal', 'ano','cada',
             'aroldo', 'raimundo', 'cedraz', 'tiago', 'césar', 'cpf', 'oliveira', 'paulo', 'astra', '(art.', 'pereira',
             'ltda.', '13.', 'iii,', 'us$', '(fl.','20):', '2º', '2º', 'spe', 'fl.', '(ato', 'nº);']

blacklist_partes = ['_','–',' ','Augusto','augusto','Carlos' "$",")","(",'/2015',"walton",'alencar','alves',".","(...)","[...]",
             'sobre', 'sr.', 'art.', 'que', 'ser', 'ii', 'ii,', 'inciso,', 'de$', 'nº', '(peça','§','(cpf','tc','maria',
             'josé','costa','que,','quanto','(',')','75.','19.','-.','(peças-11),','silva', 'tcu', 'ainda','além', 'qualquer',
             'parte','partes','item','tais','dias','tal', 'ano']



def remove_blacklist(corpus, blacklist):
    """
    removes a previously defined blacklist from the corpus, and also removes punctuation and shortwords
    """
    shortword = re.compile(r'\W*\b\w{1,1}\b') # regex to remove word up to 1 chars (inclusive)
    stoplist = stopwords.words('portuguese') + list(punctuation) + blacklist
    stoplist.remove('não')
    stoplist.remove('sem')
    text_total = [[i for i in str(shortword.sub('', str(j))).lower().split() if i not in stoplist] for j in corpus]
    return text_total


def filter_text_chunks(text_total, topic='PARTES'):
    """
    Gets text within the selected topic: PARTES, NEPOTISMO or CONFLITO DE INTERESSES.
    Optionally, can also be applyed the user defined list of terms about 'PARTES' (based in domain knowledge), in order to
    explicit consider chunks of text that include them.
    """
    filt_part1 = ['partes_relacionadas']
    filt_part2 = ['parte_relacionada']
    filt_part_add = ['acionistas', 'arm´s length', 'beneficiar', 'coligadas', 'condições de mercado', 'condições diferentes',
                     'condições diferentes das praticadas no mercado', 'conflito de interesse', 'controladas', 'contratação ',
                     'contratações', 'contratos', 'contratos ', 'diferentes das praticadas no mercado', 'dirigente', 'empresas subsidárias',
                     'favorecer', 'favorecimento', 'favorecimento indevido', 'fora das condições de mercado', 'fornecedores',
                     'governança', 'governança corporativa', 'grau de parentesco', 'influenciar', 'notas explicativas', 'obras',
                     'participação', 'parâmetro de mercado', 'preços superfaturados', 'relacionamentos', 'risco', 'riscos', 'serviços de engenharia',
                     'subcontratação', 'sócio', 'sócio privado', 'sócios comuns', 'sócios em comum', 'termos', 'termos e condições de mercado',
                     'transações', 'transparência ']

    filt_nep = ['nepotismo']
    filt_inter = ['conflito_de_interesses', 'conflito', 'interesse', 'interesses', 'conflito de interesse', 'conflito de interesses']
    text_term = []
    if topic == 'PARTES':
        # text_term = [j for j in text_total if all(i in j for i in filt_part1 or filt_part2) and any(i in j for i in filt_part_add)]
        text_term = [j for j in text_total if all(i in j for i in filt_part1 or filt_part2)]
    elif topic == 'NEPOTISMO':
        text_term = [j for j in text_total if all(i in j for i in filt_nep)]
    elif topic == 'CONF_INTERESSE':
        text_term = [j for j in text_total if any(i in j for i in filt_inter)]
    return text_term


def filter_corpus(text_term, text_total, which_text='TOTAL'):
    """
    defines if generated corpus is based on entire data or only on filtered terms data
    """
    if which_text == 'TERM':
        dict_term = corpora.Dictionary(text_term)
        corpus_term = [dict_term.doc2bow(i) for i in text_term]
        tfidf_term = models.TfidfModel(corpus_term)  # step 1 -- initialize a model
        corpus_tfidf = tfidf_term[corpus_term]  # step 2 -- use the model to transform vectors
    elif which_text == 'TOTAL':
        dict_total = corpora.Dictionary(text_total)
        corpus_total = [dict_total.doc2bow(i) for i in text_total]
        tfidf_total = models.TfidfModel(corpus_total)  # step 1 -- initialize a model
        corpus_tfidf = tfidf_total[corpus_total]  # step 2 -- use the model to transform vectors
    return corpus_tfidf


def gera_dicts(text_term, text_total):
    """
    generates dictionaries with text and total corpus, for posterior analysis
    """
    dict_term = corpora.Dictionary(text_term)
    dict_total = corpora.Dictionary(text_total)
    return dict_term, dict_total


def generate_lda(total_topics, text_term, text_total, which_text='TOTAL'):
    """
    LDA corpus generation based on choice of entire corpus data or filtered corpus data (paragraph)
    """
    if which_text == 'TERM':
        dict_term = corpora.Dictionary(text_term)
        corpus_term = [dict_term.doc2bow(i) for i in text_term]
        lda_term = models.LdaModel(corpus_term, id2word=dict_term, num_topics=total_topics)
        corpus_lda_term = lda_term[corpus_term]  # create a double wrapper over the original corpus: bow->tfidf->fold-in-lsi\
        out = [lda_term, corpus_lda_term]
    elif which_text == 'TOTAL':
        dict_total = corpora.Dictionary(text_total)
        corpus_total = [dict_total.doc2bow(i) for i in text_total]
        lda_total = models.LdaModel(corpus_total, id2word=dict_total, num_topics=total_topics)
        corpus_lda_total = lda_total[corpus_total]  # create a double wrapper over the original corpus: bow->tfidf->fold-in-lsi
        out = [lda_total, corpus_lda_total]
    return out


def show_N_terms(lda_term, lda_total, total_topics):
    """
    show terms by topic, in relevance order
    """
    # Show important words in the topics:
    print("Term: ", lda_term.show_topics(total_topics))
    print("-----")
    print("Total: ", lda_total.show_topics(total_topics))
    return


def prepare_data_clustermap(lda_term, lda_total, total_topics):
    """
    generates input dataframes for posterior clustermap generation
    """
    data_lda_term = {i: OrderedDict(lda_term.show_topic(i, 10)) for i in range(total_topics)}
    data_lda_total = {i: OrderedDict(lda_total.show_topic(i, 10)) for i in range(total_topics)}

    df_lda_term = pd.DataFrame(data_lda_term).fillna(0)
    df_lda_term = df_lda_term.T
    print("df_lda_term shape:", df_lda_term.shape)

    df_lda_total = pd.DataFrame(data_lda_total).fillna(0)
    df_lda_total = df_lda_total.T
    print("df_lda_total shape:", df_lda_total.shape)
    return df_lda_term, df_lda_total


def generate_clustermap(df_lda_total, df_lda_term, which_text='TOTAL'):
    """
    clustermap generation, based on entire data or on filtered corpus data (paragraph)
    """
    if which_text=='TERM':
        g = sns.clustermap(df_lda_term.corr(), standard_scale=1, center=0, cmap="RdBu",
                          metric='cosine', linewidths=.05, figsize=(15, 15))
        plt.setp(g.ax_heatmap.yaxis.get_majorticklabels(), rotation=0)
        # plt.savefig('clustermap_term.jpg')
        plt.show()

    if which_text=='TOTAL':
        g = sns.clustermap(df_lda_total.corr(), standard_scale=1, center=0, cmap="RdBu",
                           metric='cosine', linewidths=.05, figsize=(15, 15))
        plt.setp(g.ax_heatmap.yaxis.get_majorticklabels(), rotation=0)
        # plt.savefig('clustermap_total.jpg')
        plt.show()
    return


def generate_lda_vis(lda_term, corpus_lda_term, dict_term):
    """
    Intertopic distance generation (paragraph)
    """
    pyLDAvis.enable_notebook()
    panel = pyLDAvis.gensim.prepare(lda_term, corpus_lda_term, dict_term, mds='PCoA')
    return panel
